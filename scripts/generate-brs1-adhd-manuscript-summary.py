#!/usr/bin/env python3
"""Generate BRS1 structure + ADHD studies manuscript summary as Word document."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

OUTPUT = "/Users/paulhouston/Food Science/food-science/manuscript/BRS1-ADHD-structure-and-evidence-summary.docx"

TITLE = "BRS1 — Neurotransmitter Regulation: Framework Structure and ADHD-Relevant Evidence"

BODY = (
    "Within the BRAIN Framework, BRS1 (Neurotransmitter Regulation) describes the signal layer "
    "supporting attention, motivation, mood, and behavioural control. Four Functional Mechanisms (FMs) "
    "integrate diet-actionable Primary Mechanisms (PMs): BRS1(FM1) Monoaminergic Function (PM1–PM4: amino-acid "
    "availability, LAT1 transport, noradrenergic executive modulation, serotonergic regulation); BRS1(FM2) "
    "Cholinergic Function (PM5: acetylcholine synthesis); BRS1(FM3) Membrane Lipid Integrity (PM6: neuronal DHA "
    "incorporation); and BRS1(FM4) Excitatory–Inhibitory Balance (PM7–PM10: GABA–glutamate balance, GABA synthesis, "
    "glutamate clearance, excitotoxicity modulation). Key Constraints (shared substrate pools) and Specific Mechanisms "
    "(phenotype, genetic, and cross-system interpretation layers) provide additional structural context.\n\n"
    "In ADHD, neurobiology spans multiple transmitter systems. Dopaminergic alterations vary by subtype, stage, and "
    "brain region without supporting a simple global hypo-dopaminergic model (MacDonald et al., 2024); dopamine is "
    "linked to attention network switching (Santos et al., 2019). Meal-level amino-acid sufficiency and tryptophan/tyrosine "
    "biology connect dietary protein to catecholaminergic and serotonergic precursor supply (Wang et al., 2019; Aquili, 2020; "
    "Reimherr & Ward, 1987). Noradrenergic signalling relates to attention regulation (O'Donnell et al., 2012); serotonin "
    "to affective and impulsive phenotypes, including emotion dysregulation co-occurring with ADHD (Oades, 2010; Banerjee "
    "& Nandagopal, 2015; Shaw et al., 2014). Cholinergic evidence includes low choline intakes and altered muscarinic "
    "receptor binding (Derbyshire et al., 2023; Johansson et al., 2013). Omega-3 membrane biology is reviewed in ADHD "
    "contexts (Chang, 2021; McNamara & Carlson, 2006), with supplementation trials reported (Huss et al., 2010). "
    "Excitation–inhibition imbalance is supported by reduced GABA and altered glutamate in ADHD cohorts (Edden et al., 2012; "
    "Puts et al., 2020; Maltezos et al., 2014)."
)

REFERENCES = [
    "Aquili, L. (2020). The role of tryptophan and tyrosine in executive function and reward processing. International Journal of Tryptophan Research, 13. https://doi.org/10.1177/1178646920964825",
    "Banerjee, E., & Nandagopal, K. (2015). Does serotonin deficit mediate susceptibility to ADHD? Neurochemistry International, 82, 52–68. https://doi.org/10.1016/j.neuint.2015.02.001",
    "Chang, J. P.-C. (2021). Personalised medicine in child and adolescent psychiatry: Focus on omega-3 polyunsaturated fatty acids and ADHD. Brain, Behavior, & Immunity—Health, 16, 100310. https://doi.org/10.1016/j.bbih.2021.100310",
    "Derbyshire, E., & Maes, M. (2023). The role of choline in neurodevelopmental disorders—A narrative review focusing on ASC, ADHD and dyslexia. Nutrients, 15(13), 2876. https://doi.org/10.3390/nu15132876",
    "Edden, R. A. E., Crocetti, D., Zhu, H., Gilbert, D. L., & Mostofsky, S. H. (2012). Reduced GABA concentration in attention-deficit/hyperactivity disorder. Archives of General Psychiatry, 69(7). https://doi.org/10.1001/archgenpsychiatry.2011.2280",
    "Huss, M., Völp, A., & Stauss-Grabo, M. (2010). Supplementation of polyunsaturated fatty acids, magnesium and zinc in children seeking medical advice for attention-deficit/hyperactivity problems—An observational cohort study. Lipids in Health and Disease, 9, 105. https://doi.org/10.1186/1476-511X-9-105",
    "Johansson, J., Landgren, M., Fernell, E., Lewander, T., & Venizelos, N. (2013). Decreased binding capacity (Bmax) of muscarinic acetylcholine receptors in fibroblasts from boys with attention-deficit/hyperactivity disorder (ADHD). ADHD Attention Deficit and Hyperactivity Disorders, 5(3), 267–271. https://doi.org/10.1007/s12402-013-0103-0",
    "MacDonald, H. J., Kleppe, R., Szigetvari, P. D., & Haavik, J. (2024). The dopamine hypothesis for ADHD: An evaluation of evidence accumulated from human studies and animal models. Frontiers in Psychiatry, 15. https://doi.org/10.3389/fpsyt.2024.1492126",
    "Maltezos, S., Horder, J., Coghlan, S., et al. (2014). Glutamate/glutamine and neuronal integrity in adults with ADHD: A proton MRS study. Translational Psychiatry, 4(3), e373. https://doi.org/10.1038/tp.2014.11",
    "McNamara, R. K., & Carlson, S. E. (2006). Role of omega-3 fatty acids in brain development and function: Potential implications for the pathogenesis and prevention of psychopathology. Prostaglandins, Leukotrienes and Essential Fatty Acids, 75(4–5), 329–349. https://doi.org/10.1016/j.plefa.2006.07.010",
    "Oades, R. D. (2010). The role of serotonin in attention-deficit hyperactivity disorder (ADHD). In Progress in Brain Research (Vol. 180, pp. 565–584). https://doi.org/10.1016/S1569-7339(10)70101-6",
    "O'Donnell, J., Zeppenfeld, D., McConnell, E., Pena, S., & Nedergaard, M. (2012). Norepinephrine: A neuromodulator that boosts the function of multiple cell types to optimize CNS performance. Neurochemical Research, 37(11), 2496–2512. https://doi.org/10.1007/s11064-012-0818-x",
    "Puts, N. A. J., Ryan, M., Oeltzschner, G., et al. (2020). Reduced striatal GABA in unmedicated children with ADHD at 7T. Psychiatry Research: Neuroimaging, 301, 111082. https://doi.org/10.1016/j.pscychresns.2020.111082",
    "Reimherr, F. W., & Ward, M. F. (1987). An open trial of L-tyrosine in the treatment of attention deficit disorder, residual type. American Journal of Psychiatry, 144(8), 1071–1073. https://doi.org/10.1176/ajp.144.8.1071",
    "Santos, P. H., Gonçalves, R., & Pedroso, S. (2019). Methylphenidate and default-mode network activation: Systematic review. Revista de Neurología, 68(10), 417. https://doi.org/10.33588/rn.6810.2018487",
    "Shaw, P., Stringaris, A., Nigg, J., & Leibenluft, E. (2014). Emotion dysregulation in attention deficit hyperactivity disorder. American Journal of Psychiatry, 171(3), 276–293. https://doi.org/10.1176/appi.ajp.2013.13070966",
    "Wang, L.-J., Yu, Y.-H., Fu, M.-L., et al. (2019). Dietary profiles, nutritional biochemistry status, and attention-deficit/hyperactivity disorder: Path analysis for a case-control study. Journal of Clinical Medicine, 8(5), 709. https://doi.org/10.3390/jcm8050709",
]


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading(TITLE, level=1)
    title.runs[0].font.size = Pt(14)

    for para_text in BODY.split("\n\n"):
        p = doc.add_paragraph(para_text)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    doc.add_heading("References", level=2)
    for ref in REFERENCES:
        p = doc.add_paragraph(ref, style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    import os
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)

    word_count = len(BODY.split())
    print(f"Saved: {OUTPUT}")
    print(f"Body word count: {word_count}")


if __name__ == "__main__":
    main()
