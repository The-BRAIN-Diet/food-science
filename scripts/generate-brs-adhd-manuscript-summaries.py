#!/usr/bin/env python3
"""Generate BRS1–BRS6 framework-paper summaries for the Miguel manuscript.

Schema: system/brs-adhd-manuscript-summary-schema.md (v3.0)
Objective: justify each BRS as a distinct adaptive regulatory system in ADHD —
framework architecture, not literature review.
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

ROOT = Path(__file__).resolve().parent.parent

OUTPUT = (
    "/Users/paulhouston/Food Science/food-science/manuscript/"
    "BRS1-6-ADHD-structure-and-evidence-summary.docx"
)

DOCUMENT_INTRO = (
    "Prepared for Miguel. This document supports the BRAIN Framework paper — a "
    "framework manuscript that introduces Biological Regulatory System architecture "
    "and its alignment with adaptive regulation and allostasis.\n\n"
    "Section 1 states the conceptual contribution: why the framework has the "
    "architecture it does, and how allostatic theory provides the most appropriate "
    "interpretive lens. Sections 2–7 justify why each of the six principal "
    "Biological Regulatory Systems deserves to exist within that architecture.\n\n"
    "These summaries explain architecture, not the knowledge base. Detailed "
    "Functional Mechanisms, Cross-BRS Dependencies, Key Constraints and Primary "
    "Mechanisms are implemented on the website and are not reproduced here.\n\n"
    "Schema: system/brs-adhd-manuscript-summary-schema.md (v3.2). Target length: "
    "approximately 400–450 words (conceptual section); approximately 700–800 words "
    "per BRS summary (canonical template applied to all six systems)."
)

ALLOSTASIS_FRAMEWORK_DISCUSSION = (
    "Most nutritional frameworks are organised around nutrients, pathways, "
    "diseases, or organs. The BRAIN Framework proposes a different architecture: "
    "nutrition organised around adaptive regulatory capacities within "
    "interconnected Biological Regulatory Systems whose collective performance "
    "determines adaptive resilience. The framework was not derived from "
    "allostatic theory. It emerged independently through the systematic "
    "organisation of nutritional systems biology, and its subsequent convergence "
    "with the principles articulated by Sterling & Eyer (1988), McEwen & "
    "Wingfield (2003) and McEwen (2006) represents an important scientific "
    "observation rather than a retrospective justification. Allostatic theory "
    "explains the biology of adaptive regulation; it provides the most appropriate "
    "conceptual lens through which to interpret the framework, not its historical "
    "origin.\n\n"
    "The framework does not treat neurotransmission, inflammation, mitochondrial "
    "function, gut ecology and metabolic regulation as independent targets. It "
    "views them as mutually constraining regulatory systems in which impairment "
    "in one Biological Regulatory System progressively limits performance in "
    "others. BRS1 supports neurochemical communication; BRS2 sustains "
    "biosynthetic and repair capacity; BRS3 preserves immune and redox control; "
    "BRS4 governs cellular energy production and metabolic signalling; BRS5 "
    "serves as the principal biological interface through which dietary and "
    "microbial signals enter the adaptive regulatory network; and BRS6 "
    "coordinates resource allocation across the network. Cross-BRS dependencies "
    "— BRS6 constrains BRS4; BRS4 constrains BRS1; BRS3 constrains BRS1; BRS5 "
    "constrains BRS3; BRS2 constrains BRS1 — transform adaptive regulation from "
    "a conceptual idea into an explicit biological architecture.\n\n"
    "The framework operationalises allostatic principles as a translational "
    "nutritional interface. Adaptive regulatory capacities are specified through "
    "defined Biological Regulatory Systems, Functional Mechanisms and Primary "
    "Mechanisms; cross-system relationships are made explicit; phenomes provide a "
    "clinically interpretable bridge between underlying biological regulation and "
    "observable cognitive, emotional and behavioural outcomes; and nutritional "
    "and lifestyle interventions are linked to defined intervention points. By "
    "rendering biological assumptions and cross-system dependencies explicit, the "
    "architecture converts abstract systems theory into biologically plausible "
    "predictions intended to be empirically tested, refined or rejected as "
    "evidence accumulates. Its contribution is therefore organisational and "
    "methodological: one structured, scientifically falsifiable architecture "
    "through which nutritional science can formulate and investigate hypotheses "
    "about adaptive regulation.\n\n"
    "Picard et al. (2018) and Picard & Shirihai (2022) illustrate how such "
    "coordination may operate at the cellular level. Their mitochondrial "
    "information processing system proposes that mitochondria continuously sense, "
    "integrate and respond to neuroendocrine, metabolic and environmental "
    "signals — providing a mechanistic basis for the energetic and signalling "
    "exchange through which cross-BRS dependencies are coordinated.\n\n"
    "The principal contribution of this paper is not a new theory of allostasis. "
    "It is one translational nutritional architecture — organised around "
    "interconnected Biological Regulatory Systems representing adaptive regulatory "
    "capacities — that aligns with allostatic theory and renders its principles "
    "operationally accessible to nutritional systems biology. Allostasis is the "
    "organising principle through which the framework's architecture is structured "
    "and interpreted; the framework is the means through which nutritional science "
    "can engage with, and empirically interrogate, the biology of adaptive "
    "regulation."
)

ALLOSTASIS_FRAMEWORK_REFERENCES = [
    "Sterling, P., & Eyer, J. (1988). Allostasis: A new paradigm to explain arousal pathology. In S. Fisher & J. Reason (Eds.), Handbook of Life Stress, Cognition and Health (pp. 629–649). Wiley.",
    "McEwen, B. S., & Wingfield, J. C. (2003). The concept of allostasis in biology and biomedicine. Hormones and Behavior, 43(1), 2–15. https://doi.org/10.1016/s0018-506x(02)00024-7",
    "McEwen, B. S. (2006). Protective and damaging effects of stress mediators in health and disease. Dialogues in Clinical Neuroscience, 8(4), 377–384. https://doi.org/10.31887/DCNS.2006.8.4/bmcewen",
    "Picard, M., McEwen, B. S., Epel, E. S., & Sandi, C. (2018). An energetic view of stress: Focus on mitochondria. Frontiers in Neuroendocrinology, 49, 72–85. https://doi.org/10.1016/j.yfrne.2018.01.001",
    "Picard, M., & Shirihai, O. S. (2022). Mitochondrial signal transduction. Cell Metabolism, 34(11), 1620–1653. https://doi.org/10.1016/j.cmet.2022.10.008",
]

SECTIONS = [
    {
        "title": "BRS1 — Neurotransmitter Regulation",
        "biological_regulatory_identity": (
            "BRS1 (Neurotransmitter Regulation) exists within the BRAIN "
            "Framework because adaptive performance in attention, motivation, "
            "arousal and behavioural control depends on sustained neurochemical "
            "communication — not on any single transmitter pathway. Dopamine, "
            "noradrenaline, serotonin, acetylcholine, GABA and glutamate "
            "function as an integrated communication network whose synthesis, "
            "transport, receptor signalling, recycling and "
            "inhibition–excitation balance must be coordinated as cognitive "
            "and behavioural demand fluctuates. Collectively, these "
            "neurotransmitter systems represent the downstream neurochemical "
            "expression of multiple interacting adaptive regulatory capacities "
            "rather than independent biological targets.\n\n"
            "The framework organises this capacity as a distinct Biological "
            "Regulatory System because pathway-level interventions cannot "
            "substitute for system-level regulation. Membrane lipid composition "
            "shapes receptor environments; cholinergic tone modulates attentional "
            "filtering alongside catecholaminergic drive; and "
            "excitation–inhibition balance constrains the bandwidth within which "
            "monoaminergic output remains proportionate. BRS1 therefore "
            "represents the adaptive regulatory capacity for coordinated "
            "neurochemical communication — a regulatory tier whose existence is "
            "architecturally necessary because no isolated neurotransmitter "
            "domain explains how the brain sustains flexible attention, emotional "
            "regulation and behavioural control under changing conditions."
        ),
        "landmark_adhd_evidence": (
            "Within the BRAIN Framework, ADHD serves as the principal "
            "translational exemplar for BRS1 — not because the disorder is "
            "reducible to neurotransmitter pathology, but because it demonstrates "
            "how disruption of integrated neurochemical communication manifests "
            "in clinically observable regulatory failure. Dopaminergic dysfunction "
            "remains an important component of ADHD biology, yet landmark human "
            "evidence does not support a simple global hypo-dopaminergic model. "
            "MacDonald et al. (2024), synthesising four decades of human and "
            "animal data, conclude that dopamine involvement is established but "
            "that a uniform dopamine-deficit account cannot explain the "
            "disorder's heterogeneity across subtypes, developmental stages and "
            "brain regions — a finding that justifies the framework's refusal "
            "to organise nutrition around a single monoaminergic target. This "
            "illustrates an important principle of the framework: clinical "
            "phenomes are more plausibly interpreted as the functional "
            "expression of multiple interacting adaptive regulatory systems "
            "than as the consequence of dysfunction within a single "
            "neurotransmitter pathway.\n\n"
            "Parallel ADHD cohort evidence extends this argument across the "
            "integrated communication network. Reduced cortical and striatal GABA "
            "in medication-naïve individuals (Edden et al., 2012; Puts et al., "
            "2020) indicates constrained excitation–inhibition balance; "
            "glutamate–glutamine disturbances correlate with attentional "
            "impairment (Maltezos et al., 2014); serotonergic biology "
            "contributes to emotional dysregulation that frequently co-occurs "
            "with ADHD symptom patterns (Oades, 2010; Banerjee & Nandagopal, "
            "2015); and cholinergic receptor alterations in neurodevelopmental "
            "contexts (Johansson et al., 2013) implicate acetylcholine alongside "
            "catecholamines in attentional modulation. ADHD is thus not "
            "interpreted within the framework as a disorder of a single "
            "neurotransmitter. Rather, altered neurochemical communication "
            "represents the downstream phenotypic expression of multiple "
            "interacting adaptive regulatory systems whose collective performance "
            "determines the capacity to sustain attention, motivation, emotional "
            "regulation and behavioural control.\n\n"
            "Meal-level amino-acid sufficiency studies (Wang et al., 2019) "
            "further illustrate the translational logic: dietary substrates "
            "influence precursor pools, yet their clinical significance depends "
            "on how upstream bioenergetic, inflammatory and neuroendocrine "
            "constraints permit those precursors to sustain balanced signalling. "
            "The ADHD exemplar therefore answers the architectural question a "
            "framework reviewer is likely to pose — whether this organisation "
            "explains anything useful — by demonstrating that an integrated "
            "neurochemical communication model accommodates the breadth of "
            "established ADHD biology more coherently than dopamine-centric "
            "accounts."
        ),
        "systems_integration": (
            "BRS1 occupies a structurally central position within the "
            "integrated BRS network as the principal downstream expression "
            "layer of adaptive regulation — the system whose performance is "
            "most immediately visible in attention and behaviour, yet rarely "
            "limited by neurochemical biology alone. Its synthesis, signalling "
            "and recycling depend on upstream Biological Regulatory Systems: "
            "amino-acid availability and competitive precursor transport within "
            "BRS1 itself; methyl-donor and tetrahydrobiopterin-dependent "
            "cofactor chemistry from BRS2 (Fanet et al., 2021); inflammatory "
            "and kynurenine-mediated reshaping of the neurochemical environment "
            "from BRS3 (Savitz, 2019; Slavich & Irwin, 2014); mitochondrial "
            "bioenergetic supply that sets the energetic ceiling for sustained "
            "synaptic activity from BRS4 (Harris et al., 2012; Picard et al., "
            "2018); gut-derived vagal and metabolite signalling from BRS5; "
            "and neuroendocrine resource allocation from BRS6 (McEwen, 2006).\n\n"
            "The framework makes these dependencies explicit as mutual "
            "constraints: constrained bioenergetics, inflammatory load or "
            "cofactor insufficiency progressively limit neurochemical performance "
            "even when precursor supply appears adequate. Altered neurotransmitter "
            "regulation should therefore be interpreted as the downstream "
            "expression of interacting adaptive regulatory capacities, not the "
            "consequence of a single neurotransmitter deficit. Landmark "
            "architectural evidence supports this organisation rather than "
            "reviewing the field: Harris et al. (2012) establish that synaptic "
            "signalling consumes the largest fraction of neuronal energy "
            "expenditure, positioning BRS4 as an upstream enabler of BRS1; "
            "Fanet et al. (2021) link one-carbon and cofactor chemistry to "
            "monoamine biosynthesis, supporting the BRS2→BRS1 dependency; and "
            "Picard et al. (2018) reframe mitochondria as dynamic integrators "
            "of neuroendocrine and metabolic signals, providing a cellular basis "
            "for cross-BRS energetic and signalling exchange."
        ),
        "role_within_adaptive_network": (
            "BRS1 represents the neurochemical expression layer of the "
            "framework. It translates the performance of upstream Biological "
            "Regulatory Systems into coordinated neurotransmitter signalling "
            "that becomes observable as attention, motivation, emotional "
            "regulation and behavioural control. Because it is constrained by "
            "multiple upstream systems, altered neurotransmission is interpreted "
            "as an indicator of wider adaptive regulatory performance rather "
            "than an isolated biological deficit. Protein quality, precursor "
            "availability, membrane lipids and dietary cofactors therefore "
            "become biologically meaningful not because they target individual "
            "neurotransmitters, but because they support the wider adaptive "
            "regulatory capacity for coordinated neurochemical communication."
        ),
        "references": {
            "adhd_landmark": [
                "MacDonald, H. J., Kleppe, R., Szigetvari, P. D., & Haavik, J. (2024). The dopamine hypothesis for ADHD: An evaluation of evidence accumulated from human studies and animal models. Frontiers in Psychiatry, 15, 1492126. https://doi.org/10.3389/fpsyt.2024.1492126",
                "Edden, R. A. E., Crocetti, D., Zhu, H., Gilbert, D. L., & Mostofsky, S. H. (2012). Reduced GABA concentration in attention-deficit/hyperactivity disorder. Archives of General Psychiatry, 69(7), 750–753. https://doi.org/10.1001/archgenpsychiatry.2011.2280",
                "Puts, N. A. J., Ryan, M., Oeltzschner, G., et al. (2020). Reduced striatal GABA in unmedicated children with ADHD at 7T. Psychiatry Research: Neuroimaging, 301, 111082. https://doi.org/10.1016/j.pscychresns.2020.111082",
                "Maltezos, S., et al. (2014). Glutamate/glutamine and neuroradiologic markers of the brain in adults with ADHD. Psychiatry Research: Neuroimaging, 223(2), 121–128. https://doi.org/10.1016/j.pscychresns.2014.05.010",
                "Oades, R. D. (2010). The role of serotonin in ADHD. In C. Stanford & R. Tannock (Eds.), Behavioral Neuroscience of Attention Deficit Hyperactivity Disorder and Its Treatment (pp. 65–84). Springer. https://doi.org/10.1007/7854_2010_112",
                "Banerjee, E., & Nandagopal, K. (2015). Does serotonin deficit explain the core symptoms of ADHD? Journal of Attention Disorders, 19(1), 3–11. https://doi.org/10.1177/1087054712460387",
                "Johansson, M., et al. (2013). Decreased mAChR binding in neurodevelopmental disorders. NeuroImage, 66, 233–239. https://doi.org/10.1016/j.neuroimage.2012.10.068",
                "Wang, L.-J., et al. (2019). Dietary profiles, plasma levels of amino acids, and the risk of attention deficit hyperactivity disorder in children. Nutrients, 11(9), 2088. https://doi.org/10.3390/nu11092088",
            ],
            "systems_integration": [
                "Fanet, H., Capuron, L., Castanon, N., Calon, F., & Vancassel, S. (2021). Tetrahydrobiopterin (BH4) pathway: From metabolism to neuropsychiatry. Current Neuropharmacology, 19(5), 591–609. https://doi.org/10.2174/1570159X18666200729103529",
                "Savitz, J. (2019). The kynurenine pathway: A finger in every pie. Molecular Psychiatry, 25(1), 131–147. https://doi.org/10.1038/s41380-019-0414-4",
                "Slavich, G. M., & Irwin, M. R. (2014). From stress to inflammation and major depressive disorder: A social signal transduction theory of depression. Psychological Bulletin, 140(3), 774–815. https://doi.org/10.1037/a0035302",
                "Harris, J. J., Jolivet, R., & Attwell, D. (2012). Synaptic energy use and supply. Neuron, 75(5), 762–777. https://doi.org/10.1016/j.neuron.2012.08.019",
                "Picard, M., McEwen, B. S., Epel, E. S., & Sandi, C. (2018). An energetic view of stress: Focus on mitochondria. Frontiers in Neuroendocrinology, 49, 72–85. https://doi.org/10.1016/j.yfrne.2018.01.001",
                "McEwen, B. S. (2006). Protective and damaging effects of stress mediators in health and disease. Dialogues in Clinical Neuroscience, 8(4), 377–384. https://doi.org/10.31887/DCNS.2006.8.4/bmcewen",
            ],
            "allostasis_and_resilience": [],
        },
    },
    {
        "title": "BRS2 — Methylation & One-Carbon Metabolism",
        "biological_regulatory_identity": (
            "Biological adaptation requires continuous methyl-group transfer, "
            "homocysteine handling, transsulfuration-linked redox chemistry and "
            "phospholipid formation — processes that share a common one-carbon "
            "metabolic infrastructure. BRS2 (Methylation & One-Carbon Metabolism) "
            "represents the integrated capacity to maintain this biochemical "
            "maintenance layer: remethylation flux, transsulfuration coupling and "
            "membrane biosynthesis operating as a coupled regulatory network "
            "rather than isolated vitamin pathways. Functional Mechanisms span "
            "homocysteine remethylation and methionine-cycle flux, "
            "transsulfuration and glutathione precursor supply, and "
            "phosphatidylcholine formation linking methyl-donor chemistry to "
            "membrane integrity and cholinergic support.\n\n"
            "The system warrants Biological Regulatory System status because "
            "constrained one-carbon throughput simultaneously limits "
            "neurotransmitter cofactor availability, glutathione-centred redox "
            "defence and membrane integrity. A disturbance in folate or B12 "
            "handling does not produce a single biochemical readout — it "
            "propagates across downstream regulatory domains. In ADHD, where "
            "cohort evidence repeatedly identifies one-carbon insufficiency, "
            "homocysteine elevation and genetic susceptibility in folate cycling, "
            "this integrated maintenance capacity is biologically relevant as "
            "shared infrastructure on which neurotransmission, redox defence and "
            "membrane biology depend."
        ),
        "landmark_adhd_evidence": (
            "Within the BRAIN Framework, ADHD serves as the principal "
            "translational exemplar for BRS2 — not because the disorder is "
            "reducible to vitamin deficiency, but because it demonstrates how "
            "impaired one-carbon maintenance propagates across multiple adaptive "
            "capacities. Meta-analytic evidence establishes measurable methyl-donor "
            "insufficiency in ADHD populations. Razavinia et al. (2024) reported "
            "lower folate and vitamin B12 status in children with ADHD compared "
            "with controls across pooled studies — supporting the view that "
            "methyl-donor availability is not merely a nutritional footnote but "
            "a reproducible biological feature of the disorder.\n\n"
            "Complementary cohort biomarker work links this insufficiency to "
            "impaired metabolic flux. Lukovac et al. (2024) found elevated "
            "homocysteine alongside disturbances in B12, vitamin D and iron "
            "markers in paediatric ADHD — a pattern consistent with constrained "
            "remethylation and downstream transsulfuration rather than isolated "
            "nutrient deficiency. Genetic meta-analysis further supports "
            "architectural relevance: Meng et al. (2022) reported association "
            "between the MTHFR 1298A>C polymorphism and ADHD susceptibility, "
            "implicating folate-cycle efficiency as a modifiable biological "
            "context rather than a deterministic genetic cause.\n\n"
            "Dietary pattern evidence extends the argument beyond biomarkers. "
            "Wang et al. (2019) linked unhealthy dietary patterns to lower B "
            "vitamin status and ADHD in case–control path analysis, while "
            "patterns rich in folate and fibre have been associated with reduced "
            "symptom burden. One-carbon biology also bridges membrane and "
            "cholinergic domains: choline insufficiency is implicated across "
            "neurodevelopmental disorders including ADHD (Derbyshire & Maes, "
            "2023), and reduced muscarinic receptor binding in boys with ADHD "
            "(Johansson et al., 2013) links methyl-donor-dependent "
            "phosphatidylcholine chemistry to attentional modulation. ADHD is "
            "therefore not interpreted as a methylation disorder alone. Rather, "
            "one-carbon insufficiency represents upstream biochemical "
            "infrastructure whose impairment may constrain neurotransmission, "
            "redox defence and membrane biology simultaneously — a pattern the "
            "framework organises as a distinct Biological Regulatory System "
            "rather than a catalogue of micronutrient deficits."
        ),
        "systems_integration": (
            "BRS2 functions as upstream biochemical infrastructure for "
            "neurotransmitter regulation and redox defence. Cofactor chemistry "
            "governed by one-carbon flux — including tetrahydrobiopterin-dependent "
            "monoamine synthesis — directly enables BRS1 performance (Fanet et "
            "al., 2021). Transsulfuration routing from homocysteine to cysteine "
            "supplies glutathione precursors that sustain antioxidant capacity "
            "within BRS3 (Kumar et al., 2017). Phosphatidylcholine formation "
            "through SAMe-dependent PEMT activity links methyl-donor throughput "
            "to membrane phospholipid renewal and long-chain omega-3 handling "
            "that supports BRS1 membrane biology.\n\n"
            "Because these outputs derive from a shared metabolic node, "
            "one-carbon insufficiency rarely announces itself as a methylation "
            "defect alone — it propagates into neurotransmission and oxidative "
            "resilience downstream. The framework makes this dependency explicit: "
            "BRS2→BRS1 and BRS2→BRS3 represent mutual constraints in which "
            "constrained remethylation or transsulfuration progressively limits "
            "cofactor supply and glutathione precursor availability even when "
            "dietary protein appears adequate. BRS2 therefore occupies a "
            "distinct position as the biochemical maintenance layer that "
            "connected regulatory systems draw upon under sustained demand."
        ),
        "role_within_adaptive_network": (
            "BRS2 provides the biochemical maintenance infrastructure shared "
            "by multiple downstream systems, coupling one-carbon metabolism to "
            "neurotransmission, membrane biology and antioxidant defence. "
            "Folate, B12, choline and methyl-donor availability from diet "
            "therefore matter not as isolated micronutrient corrections, but "
            "as inputs into the shared biochemical maintenance infrastructure "
            "on which connected regulatory systems depend. When one-carbon "
            "throughput is constrained, dietary correction of a single vitamin "
            "cannot restore downstream neurotransmitter cofactor supply, "
            "glutathione formation and membrane renewal simultaneously — "
            "justifying BRS2 as a system-level maintenance tier rather than a "
            "micronutrient checklist within the framework."
        ),
        "references": {
            "adhd_landmark": [
                "Razavinia, F., et al. (2024). Vitamins B9 and B12 in children with attention deficit hyperactivity disorder (ADHD). International Journal for Vitamin and Nutrition Research, 94(5–6), 476–484. https://doi.org/10.1024/0300-9831/a000809",
                "Lukovac, T., et al. (2024). Serum biomarker analysis in pediatric ADHD. Children, 11(4), 497. https://doi.org/10.3390/children11040497",
                "Meng, X., et al. (2022). Association between MTHFR polymorphisms and ADHD: A meta-analysis. Frontiers in Pediatrics, 10, 893789. https://doi.org/10.3389/fped.2022.893789",
                "Wang, L.-J., et al. (2019). Dietary profiles, plasma levels of amino acids, and the risk of attention deficit hyperactivity disorder in children. Nutrients, 11(9), 2088. https://doi.org/10.3390/nu11092088",
                "Derbyshire, E., & Maes, M. (2023). The role of choline in neurodevelopmental disorders. Nutrients, 15(13), 2876. https://doi.org/10.3390/nu15132876",
                "Johansson, M., et al. (2013). Decreased mAChR binding in neurodevelopmental disorders. NeuroImage, 66, 233–239. https://doi.org/10.1016/j.neuroimage.2012.10.068",
            ],
            "systems_integration": [
                "Fanet, H., Capuron, L., Castanon, N., Calon, F., & Vancassel, S. (2021). Tetrahydrobiopterin (BH4) pathway: From metabolism to neuropsychiatry. Current Neuropharmacology, 19(5), 591–609. https://doi.org/10.2174/1570159X18666200729103529",
                "Kumar, N., Kaur, G., & Kaur, S. (2017). The transsulfuration pathway. World Journal of Hepatology, 9(16), 745–750. https://doi.org/10.4254/wjh.v9.i16.745",
            ],
            "allostasis_and_resilience": [],
        },
    },
    {
        "title": "BRS3 — Inflammation & Oxidative Stress",
        "biological_regulatory_identity": (
            "Adaptive regulation requires the capacity to mount proportionate "
            "inflammatory responses, maintain antioxidant defence and resolve "
            "immune activation without chronic low-grade burden. BRS3 "
            "(Inflammation & Oxidative Stress) represents this integrated "
            "immune–redox regulatory capacity — anti-inflammatory signalling "
            "tone, antioxidant defence, gut-derived inflammatory modulation "
            "and active inflammation resolution operating as a coordinated "
            "system rather than isolated biomarkers.\n\n"
            "Inflammatory and oxidative load never remains confined to immune "
            "biology. Elevated cytokine tone, lipid peroxidation and impaired "
            "resolution reshape receptor environments, mitochondrial efficiency "
            "and neurotransmitter turnover. The system is organised as a "
            "Biological Regulatory System because immune and redox biology "
            "jointly determine the conditions under which cognition, mood and "
            "behavioural control are sustained. In ADHD, where cohort evidence "
            "repeatedly identifies elevated inflammatory markers, oxidative "
            "burden and gut-immune overlap, this integrated capacity is "
            "biologically relevant as the environmental layer within which "
            "neurotransmission and bioenergetics must operate."
        ),
        "landmark_adhd_evidence": (
            "Within the BRAIN Framework, ADHD serves as the principal "
            "translational exemplar for BRS3 — not because the disorder is "
            "reducible to inflammation, but because it demonstrates how "
            "immune–redox dysregulation constrains the wider adaptive network. "
            "Landmark ADHD cohort work links neuroimmune activation to the "
            "disorder. Chang et al. (2020), studying children and adolescents "
            "with ADHD in Taiwan, reported associations between elevated reactive "
            "oxygen species, astrocyte and microglial activation, and "
            "pro-inflammatory cytokines including IL-6, IL-1β and TNF-α — "
            "published in Brain, Behavior, and Immunity. This study anchors "
            "neuroimmune inflammatory activation as reproducible ADHD-associated "
            "biology rather than a speculative comorbidity.\n\n"
            "Independent evidence confirms elevated net oxidative burden across "
            "life stages. Bulut et al. (2007) reported significantly higher "
            "malondialdehyde — a lipid peroxidation marker — in adults with "
            "ADHD versus healthy controls. Kurhan & Alp (2021) found shifted "
            "thiol/disulfide homeostasis toward oxidation and elevated urinary "
            "8-OHdG in adult ADHD, while Miniksar et al. (2023) reported higher "
            "oxidative stress indices — convergent evidence that antioxidant "
            "defence and ROS clearance are under strain rather than simply "
            "deficient. Gut-immune overlap further supports architectural "
            "breadth: reduced microbiome alpha diversity in paediatric ADHD "
            "(Prehn-Kristensen et al., 2018) links peripheral ecological "
            "disturbance to the inflammatory tone BRS3 must contain. Wesselink "
            "et al. (2019) report immune dysfunction and elevated IgE patterns "
            "overlapping with ADHD contexts — reinforcing that inflammatory "
            "regulation spans gut-derived, allergic and neuroimmune domains "
            "rather than a single cytokine pathway.\n\n"
            "ADHD is therefore not interpreted as an inflammatory disorder "
            "alone. Rather, elevated immune activation and oxidative burden "
            "represent upstream constraints on the biological environment "
            "within which attention, mood and behavioural control are sustained. "
            "Dietary antioxidant and anti-inflammatory inputs acquire "
            "translational significance only when interpreted through this "
            "integrated immune–redox architecture — not as isolated corrections "
            "of single biomarkers. Verlaet et al. (2018) provide additional "
            "translational logic by framing dietary antioxidant treatment of ADHD "
            "through immune, epigenetic and oxidative-stress regulation — "
            "supporting the architectural view that inflammatory and redox biology "
            "are modifiable system-level contexts rather than downstream "
            "consequences of neurotransmitter dysfunction alone."
        ),
        "systems_integration": (
            "BRS3 shapes the biological environment within which downstream "
            "systems operate. Chronic inflammatory activation propagates through "
            "neuroimmune pathways — including kynurenine metabolism — that "
            "directly reshape monoaminergic signalling and excitation–inhibition "
            "balance in BRS1 (Savitz, 2019). Gut-derived immune spillover from "
            "impaired barrier containment (BRS5) amplifies systemic inflammatory "
            "tone (O'Mahony et al., 2015), while stress-mediated load allocation "
            "(BRS6) activates inflammatory signalling through coordinated "
            "neuroendocrine pathways (Slavich & Irwin, 2014). One-carbon "
            "transsulfuration from BRS2 supplies glutathione precursors that "
            "sustain antioxidant recycling within BRS3, making immune–redox "
            "performance dependent on shared biochemical infrastructure "
            "upstream.\n\n"
            "The framework makes these dependencies explicit as mutual "
            "constraints: BRS3→BRS1, BRS5→BRS3 and BRS6→BRS3 mean that "
            "strain surfacing in neurotransmission or bioenergetics often "
            "originates in immune tone and redox burden that never fully "
            "resolved. Landmark architectural evidence supports this "
            "organisation: Savitz (2019) positions kynurenine metabolism as a "
            "convergent neuroimmune interface reshaping monoaminergic biology, "
            "while Slavich & Irwin (2014) formalise stress-to-inflammation "
            "signalling as a systems-level constraint on adaptive performance."
        ),
        "role_within_adaptive_network": (
            "BRS3 determines the inflammatory and redox environment within "
            "which the remaining Biological Regulatory Systems must operate. "
            "Polyphenols, antioxidant nutrients and anti-inflammatory dietary "
            "patterns become relevant not because they target single "
            "inflammatory markers, but because they help maintain the immune "
            "and redox conditions that constrain neurotransmission, "
            "bioenergetics and stress recovery downstream. Restoring "
            "proportionate inflammatory resolution and antioxidant recycling "
            "therefore supports the wider adaptive network — a systems-level "
            "interpretation the framework reserves for BRS3 rather than "
            "nutrient-specific anti-inflammatory claims."
        ),
        "references": {
            "adhd_landmark": [
                "Chang, J. P.-C., Mondelli, V., Satyanarayanan, S. K., et al. (2020). Cortisol, inflammatory biomarkers and neurotrophins in children and adolescents with ADHD in Taiwan. Brain, Behavior, and Immunity, 88, 105–113. https://doi.org/10.1016/j.bbi.2020.05.017",
                "Bulut, M., et al. (2007). Evaluation of malondialdehyde levels in adult ADHD. Psychiatry and Clinical Psychopharmacology, 17(2), 89–94. https://doi.org/10.1080/14751740701595421",
                "Kurhan, S., & Alp, H. H. (2021). Dynamic thiol/disulfide homeostasis and 8-OHdG in adult ADHD. Psychiatry and Clinical Psychopharmacology, 31(4), 394–400. https://doi.org/10.5152/pcp.2021.21078",
                "Miniksar, O. P., et al. (2023). Effect of oxidative stress on adult ADHD. Clinical Psychopharmacology and Neuroscience, 21(1), 58–64. https://doi.org/10.9758/cpn.2023.21.1.58",
                "Prehn-Kristensen, A., et al. (2018). Reduced microbiome alpha diversity in pediatric ADHD. PLOS ONE, 13(7), e0200728. https://doi.org/10.1371/journal.pone.0200728",
                "Verlaet, A. A. J., et al. (2018). Rationale for dietary antioxidant treatment of ADHD. Nutrients, 10(4), 405. https://doi.org/10.3390/nu10040405",
                "Wesselink, E., et al. (2019). Feeding immunity: Physiological and behavioral responses to food allergy in ADHD. Nutrients, 11(10), 2351. https://doi.org/10.3390/nu11102351",
            ],
            "systems_integration": [
                "Savitz, J. (2019). The kynurenine pathway: A finger in every pie. Molecular Psychiatry, 25(1), 131–147. https://doi.org/10.1038/s41380-019-0414-4",
                "Slavich, G. M., & Irwin, M. R. (2014). From stress to inflammation and major depressive disorder: A social signal transduction theory of depression. Psychological Bulletin, 140(3), 774–815. https://doi.org/10.1037/a0035302",
                "O'Mahony, S. M., Clarke, G., Borre, Y. E., Dinan, T. G., & Cryan, J. F. (2015). Serotonin, tryptophan metabolism and the brain-gut-microbiome axis. Behavioural Brain Research, 277, 32–48. https://doi.org/10.1016/j.bbr.2014.07.027",
            ],
            "allostasis_and_resilience": [],
        },
    },
    {
        "title": "BRS4 — Mitochondrial Function & Bioenergetics",
        "biological_regulatory_identity": (
            "Every adaptive process — including sustained attention, synaptic "
            "signalling and stress recovery — depends on cellular capacity to "
            "generate ATP, protect mitochondrial function and flexibly route "
            "metabolic substrates. BRS4 (Mitochondrial Function & Bioenergetics) "
            "represents this integrated bioenergetic capacity: cellular energy "
            "production, mitochondrial resilience, substrate utilisation "
            "flexibility and adaptive energetic remodelling operating as a "
            "coordinated regulatory system spanning electron transport, redox "
            "stability, fuel switching and organelle biogenesis.\n\n"
            "Bioenergetic reserve is seldom measured directly; its erosion is "
            "often recognised first through attention lapses, cognitive stamina "
            "or neurotransmitter instability. The system warrants Biological "
            "Regulatory System status because mitochondrial capacity sets the "
            "energetic ceiling beneath which every other regulatory system must "
            "operate. In ADHD, where direct patient-derived cellular evidence "
            "demonstrates impaired mitochondrial respiration and controlled "
            "bioenergetic substrate interventions modify function, BRS4 "
            "represents a biologically distinct and architecturally necessary "
            "layer of adaptive capacity."
        ),
        "landmark_adhd_evidence": (
            "Within the BRAIN Framework, ADHD serves as the principal "
            "translational exemplar for BRS4 — not because the disorder is "
            "reducible to mitochondrial disease, but because it demonstrates "
            "how bioenergetic constraint limits the performance of connected "
            "regulatory systems. Primary ADHD-derived cellular evidence "
            "establishes mitochondrial dysfunction as a feature of the disorder. "
            "Verma et al. (2016) demonstrated in cybrid cell models generated "
            "from ADHD patient platelets reduced cellular and mitochondrial "
            "respiration, lower complex V activity, membrane-potential loss and "
            "elevated oxidative stress — providing direct evidence that "
            "bioenergetic capacity is impaired at the subcellular level in ADHD.\n\n"
            "Controlled intervention evidence supports modifiability at the system "
            "level. van Oudheusden & Scholte (2002) reported that carnitine "
            "supplementation improved behavioural and functional outcomes in "
            "children with ADHD — demonstrating that a bioenergetic substrate "
            "intervention can modify ADHD-relevant function without targeting "
            "neurotransmitter pathways directly. Complementary paediatric "
            "case–control work links oxidative-stress markers to mitochondrial "
            "redox handling: Verlaet et al. (2019) reported elevated glutathione "
            "and oxidative indices consistent with compensatory antioxidant "
            "response to increased mitochondrial ROS burden. Narrative synthesis "
            "in Öğütlü et al. (2022) further consolidates ADHD-relevant "
            "mitochondrial biomarker, genetic-variation and oxidative-stress "
            "literature, supporting bioenergetics as a coherent biological domain "
            "rather than isolated trial findings.\n\n"
            "ADHD is therefore not interpreted as universal mitochondrial "
            "failure. Rather, constrained bioenergetic reserve represents an "
            "upstream ceiling on synaptic signalling, inflammatory resolution and "
            "stress recovery — capacities distributed across BRS1, BRS3 and "
            "BRS6. The exemplar answers the architectural question: when "
            "mitochondrial performance thins, downstream neurochemical and "
            "behavioural regulation falters even when precursor supply appears "
            "adequate."
        ),
        "systems_integration": (
            "BRS4 provides the energetic foundation on which neurotransmitter "
            "regulation depends. Synaptic transmission consumes the largest "
            "share of neuronal energy, and mitochondria at synapses dynamically "
            "coordinate energetic supply with signalling demand (Harris et al., "
            "2012; Picard, 2015). Gut-derived microbial metabolites (BRS5) "
            "influence substrate availability and mitochondrial signalling, while "
            "neuroendocrine and glycaemic stability (BRS6) shape mitochondrial "
            "recovery through stress-mediated metabolic allocation — a pathway "
            "formalised as mitochondrial allostatic load (Picard et al., 2014). "
            "Inflammatory and oxidative strain from BRS3 further taxes "
            "mitochondrial redox integrity, tightening the energetic ceiling "
            "under which BRS1 must sustain neurotransmitter turnover.\n\n"
            "The framework makes BRS4→BRS1, BRS5→BRS4 and BRS6→BRS4 explicit "
            "as mutual constraints: when bioenergetic reserve thins, downstream "
            "signalling and recovery falter even when neurotransmitter pathways "
            "appear biochemically intact. Landmark architectural evidence "
            "supports this organisation: Harris et al. (2012) establish synaptic "
            "energy demand as the principal neuronal ATP sink, positioning BRS4 "
            "as an upstream enabler of BRS1; Picard et al. (2014) link "
            "glucocorticoid dynamics to mitochondrial allostatic load, "
            "connecting BRS6 resource allocation to BRS4 recovery capacity. "
            "Almutairi et al. (2024) further consolidate mitochondrial "
            "bioenergetic regulation and mitophagy as adaptive remodelling "
            "processes whose impairment would be expected to tighten this "
            "energetic ceiling across the network."
        ),
        "role_within_adaptive_network": (
            "BRS4 establishes the energetic reserve that limits sustained "
            "performance across every connected Biological Regulatory System. "
            "Dietary support for mitochondrial substrate supply — including "
            "carnitine, coenzyme Q10, magnesium and patterns that sustain "
            "metabolic flexibility — acquires biological significance not as "
            "isolated energy supplementation, but as preservation of the "
            "energetic ceiling beneath which the integrated network operates. "
            "When bioenergetic reserve is thin, attentional stamina, emotional "
            "regulation and inflammatory recovery all compete for the same "
            "mitochondrial capacity — a constraint the framework makes explicit "
            "through BRS4's upstream position in the cross-BRS dependency "
            "architecture."
        ),
        "references": {
            "adhd_landmark": [
                "Verma, P., Singh, A., Nthenge-Ngumbau, D. N., et al. (2016). Attention deficit-hyperactivity disorder suffers from mitochondrial dysfunction. BBA Clinical, 6, 153–158. https://doi.org/10.1016/j.bbacli.2016.10.003",
                "van Oudheusden, L. J., & Scholte, H. R. (2002). Efficacy of carnitine in the treatment of children with ADHD. Prostaglandins, Leukotrienes and Essential Fatty Acids, 67(1), 33–38. https://doi.org/10.1054/plef.2002.0378",
                "Verlaet, A. A. J., et al. (2019). Oxidative stress and immune markers in pediatric ADHD. European Child & Adolescent Psychiatry, 28(6), 719–729. https://doi.org/10.1007/s00787-018-1240-1",
                "Öğütlü, A., et al. (2022). Mitochondrial dysfunction in ADHD: A narrative review. Psychiatry and Clinical Psychopharmacology, 32(4), 281–290. https://doi.org/10.5152/pcp.2022.22014",
            ],
            "systems_integration": [
                "Harris, J. J., Jolivet, R., & Attwell, D. (2012). Synaptic energy use and supply. Neuron, 75(5), 762–777. https://doi.org/10.1016/j.neuron.2012.08.019",
                "Picard, M., Juster, R.-P., & McEwen, B. S. (2014). Mitochondrial allostatic load puts the 'gluc' back in glucocorticoids. Nature Reviews Endocrinology, 10(5), 303–310. https://doi.org/10.1038/nrendo.2014.22",
                "Picard, M. (2015). Mitochondrial synapses: Intracellular communication and signal integration. Trends in Neurosciences, 38(8), 468–474. https://doi.org/10.1016/j.tins.2015.06.001",
                "Almutairi, M. A., et al. (2024). Mitochondrial bioenergetic regulation and mitophagy in ADHD. Biomedicines, 12(3), 512. https://doi.org/10.3390/biomedicines12030512",
            ],
            "allostasis_and_resilience": [],
        },
    },
    {
        "title": "BRS5 — Gut-Brain Axis & Enteric Nervous System",
        "biological_regulatory_identity": (
            "The gut is not merely a digestive organ but a continuous source of "
            "immune, metabolic and neuromodulatory signals that shape central "
            "regulatory biology. BRS5 (Gut-Brain Axis & Enteric Nervous System) "
            "represents the integrated peripheral capacity to maintain gut barrier "
            "selectivity, supportive microbial ecology, beneficial metabolite "
            "signalling and proportionate gut-to-brain communication through "
            "vagal and enteric pathways.\n\n"
            "These domains form a coordinated regulatory system because barrier "
            "integrity, microbial function and neuromodulatory signalling are "
            "interdependent: impaired containment amplifies immune spillover into "
            "BRS3; reduced microbial metabolite output constrains energetic and "
            "signalling inputs to BRS4; disrupted vagal tone alters stress-axis "
            "responsiveness within BRS6. In ADHD, where converging cohort evidence "
            "identifies gut dysbiosis, reduced short-chain fatty acid output, "
            "microbiome–reward circuitry coupling and modifiable probiotic "
            "responses, this integrated gut–brain capacity is biologically "
            "relevant as the principal peripheral interface through which diet "
            "and microbial ecology reshape the adaptive network."
        ),
        "landmark_adhd_evidence": (
            "Within the BRAIN Framework, ADHD serves as the principal "
            "translational exemplar for BRS5 — not because the disorder is "
            "reducible to dysbiosis, but because it demonstrates how peripheral "
            "regulatory biology propagates into central motivation, attention and "
            "immune-metabolic signalling. Landmark neuroimaging work links gut "
            "ecology directly to central reward circuitry in ADHD. Aarts et al. "
            "(2017) reported that gut microbiome composition in boys with ADHD "
            "correlated with neural reward anticipation on fMRI during a monetary "
            "incentive delay task — establishing a measurable gut–brain "
            "relationship in motivation biology relevant to impulsivity and "
            "attention.\n\n"
            "Complementary cohort evidence identifies microbial metabolite "
            "disruption. Steckler et al. (2024), in Brain, Behavior, & "
            "Immunity—Health, reported lower faecal short-chain fatty acid levels "
            "and gut dysbiosis patterns in children with ADHD versus controls. "
            "Reduced microbiome alpha diversity in paediatric ADHD "
            "(Prehn-Kristensen et al., 2018) and compositional differences in "
            "treatment-naïve children (Jiang et al., 2018) converge on ecological "
            "disturbance rather than a single pathogenic taxon. Intervention "
            "evidence supports modifiability: Wang et al. (2022) reported that "
            "open-label Bifidobacterium bifidum supplementation was associated "
            "with symptom change and altered gut microbiota in children with "
            "ADHD — demonstrating that peripheral microbial ecology can be "
            "influenced and that such change coincides with functional "
            "improvement, even if causal direction remains to be established.\n\n"
            "ADHD is therefore not interpreted as a gut disorder alone. Rather, "
            "altered gut–brain signalling represents a peripheral entry point "
            "through which dietary fibre, microbial diversity and enteric "
            "neuromodulation continuously reshape immune tone, bioenergetic "
            "substrate supply and stress-axis responsiveness across BRS3, BRS4 "
            "and BRS6 — a pattern the framework organises as a distinct "
            "Biological Regulatory System rather than a probiotic target list."
        ),
        "systems_integration": (
            "BRS5 propagates peripheral signals across multiple downstream "
            "systems. Gut-derived immune spillover from impaired barrier "
            "containment amplifies inflammatory tone within BRS3 (O'Mahony et "
            "al., 2015). Microbial metabolite output — including short-chain "
            "fatty acids — influences bioenergetic substrate availability and "
            "mitochondrial signalling in BRS4. Vagal neuromodulation shapes "
            "autonomic tone and HPA-axis responsiveness within BRS6 (Bravo et "
            "al., 2011), while gut–microbial handling of neurotransmitter "
            "precursors modulates the environment within which BRS1 operates.\n\n"
            "The framework makes BRS5→BRS1, BRS5→BRS3, BRS5→BRS4 and "
            "BRS5→BRS6 explicit as mutual constraints: peripheral dysfunction "
            "rarely produces an isolated gastrointestinal phenotype — it "
            "propagates through immune, energetic and neuroendocrine biology "
            "downstream. Landmark architectural evidence supports this "
            "organisation: Bravo et al. (2011) demonstrate vagal-dependent "
            "central GABAergic modulation from probiotic ingestion, while "
            "O'Mahony et al. (2015) formalise gut–microbial tryptophan handling "
            "as a continuous interface with central serotonergic and immune "
            "biology. BRS5 therefore occupies a distinct position as the "
            "peripheral environmental interface through which diet and "
            "microbiota continuously reshape the integrated adaptive network."
        ),
        "role_within_adaptive_network": (
            "BRS5 provides the principal environmental interface through "
            "which diet, microbiota and peripheral signalling continuously "
            "reshape the adaptive regulatory network. Fermentable fibre, diverse "
            "plant foods and dietary patterns supporting microbial diversity "
            "matter because they influence the peripheral entry point through "
            "which immune, metabolic and neuromodulatory signals propagate "
            "across BRS3, BRS4, BRS6 and BRS1. Gut–brain axis interventions "
            "therefore acquire architectural significance only when they support "
            "barrier integrity, ecological turnover and vagal–enteric "
            "communication as an integrated peripheral regulatory capacity — not "
            "when they target individual probiotic strains in isolation from the "
            "wider adaptive network."
        ),
        "references": {
            "adhd_landmark": [
                "Aarts, E., Ederveen, T. H. A., Naaijen, J., et al. (2017). Gut microbiome in ADHD and its relation to neural reward anticipation. PLOS ONE, 12(9), e0183509. https://doi.org/10.1371/journal.pone.0183509",
                "Steckler, R., Magzal, F., Kokot, M., et al. (2024). Disrupted gut harmony in ADHD. Brain, Behavior, & Immunity—Health, 40, 100829. https://doi.org/10.1016/j.bbih.2024.100829",
                "Prehn-Kristensen, A., et al. (2018). Reduced microbiome alpha diversity in pediatric ADHD. PLOS ONE, 13(7), e0200728. https://doi.org/10.1371/journal.pone.0200728",
                "Jiang, H.-Y., et al. (2018). Altered gut microbiota profiles in treatment-naïve children with ADHD. Frontiers in Cellular and Infection Microbiology, 8, 379. https://doi.org/10.3389/fcimb.2018.00379",
                "Wang, L.-J., et al. (2022). Effect of Bifidobacterium bifidum supplementation in children with ADHD. Nutrients, 14(19), 3942. https://doi.org/10.3390/nu14193942",
            ],
            "systems_integration": [
                "Bravo, J. A., Forsythe, P., Chew, M. V., et al. (2011). Ingestion of Lactobacillus strain regulates emotional behavior and central GABA receptor expression in a mouse via the vagus nerve. Proceedings of the National Academy of Sciences, 108(38), 16050–16055. https://doi.org/10.1073/pnas.1102999108",
                "O'Mahony, S. M., Clarke, G., Borre, Y. E., Dinan, T. G., & Cryan, J. F. (2015). Serotonin, tryptophan metabolism and the brain-gut-microbiome axis. Behavioural Brain Research, 277, 32–48. https://doi.org/10.1016/j.bbr.2014.07.027",
            ],
            "allostasis_and_resilience": [],
        },
    },
    {
        "title": "BRS6 — Metabolic & Neuroendocrine Regulation",
        "biological_regulatory_identity": (
            "Adaptive performance requires coordinated allocation of energy, "
            "inflammatory load and recovery capacity across the body — not merely "
            "adequate hormone levels in isolation. BRS6 (Metabolic & "
            "Neuroendocrine Regulation) represents this integrated whole-body "
            "capacity: glycaemic stability, HPA-axis cortisol rhythm, autonomic "
            "balance and proportionate stress–metabolic load allocation operating "
            "as the principal systems through which physiological demand is "
            "translated into constraints on cognitive and behavioural performance.\n\n"
            "BRS6 warrants Biological Regulatory System status because it "
            "coordinates how physiological resources are mobilised and restored "
            "under demand — determining whether cognition, mood and behavioural "
            "control remain proportionate or drift toward exhaustion. In ADHD, "
            "where evidence identifies altered cortisol rhythms, autonomic "
            "arousal differences, cerebral glucose metabolism shifts and elevated "
            "metabolic comorbidity, this regulatory tier is biologically relevant "
            "as the system's principal gateway for resource allocation across "
            "the integrated network — the architectural tier most directly "
            "aligned with allostatic theory."
        ),
        "landmark_adhd_evidence": (
            "Within the BRAIN Framework, ADHD serves as the principal "
            "translational exemplar for BRS6 — not because the disorder is "
            "reducible to endocrine pathology, but because it demonstrates how "
            "resource-allocation dysregulation constrains performance across "
            "connected regulatory systems. A systematic review with meta-analysis "
            "establishes HPA-axis dysregulation as a reproducible ADHD "
            "neuroendocrine signature. Chang et al. (2021), in Translational "
            "Psychiatry, reported altered basal and morning cortisol patterns in "
            "youths with ADHD compared with typically developing peers — anchoring "
            "stress-axis rhythm as a landmark regulatory domain within the "
            "disorder.\n\n"
            "Complementary cohort and mechanistic evidence extends this framing "
            "beyond cortisol alone. Chang et al. (2020) measured cortisol "
            "alongside inflammatory biomarkers and neurotrophins in Taiwanese "
            "children and adolescents with ADHD, linking neuroendocrine rhythm "
            "to immune activation that propagates into BRS3. Lane et al. (2010) "
            "differentiated sensory over-responsivity from ADHD using cortisol "
            "and autonomic markers — supporting autonomic balance as a distinct "
            "regulatory dimension within the same architectural tier. Early "
            "landmark PET work by Zametkin et al. (1990) demonstrated altered "
            "cerebral glucose metabolism in prefrontal and striatal regions in "
            "adults with hyperactivity of childhood onset, anchoring metabolic "
            "energy allocation as long-standing ADHD-relevant biology.\n\n"
            "Long-term metabolic burden extends the argument into adulthood. "
            "Di Girolamo et al. (2022) reported higher prevalence of metabolic "
            "syndrome and insulin resistance in adult ADHD outpatients versus "
            "expected population rates, while Marcelli et al. (2025) synthesise "
            "shared mechanisms bridging ADHD and metabolic disorders. ADHD is "
            "therefore not interpreted as a cortisol disorder alone. Rather, "
            "disturbed resource allocation represents the systemic logic through "
            "which glycaemic instability, inflammatory load and recovery "
            "capacity progressively constrain BRS4 bioenergetics and BRS1 "
            "neurochemical performance — the architectural question the "
            "framework poses and the ADHD exemplar helps answer."
        ),
        "systems_integration": (
            "BRS6 coordinates resource allocation across the integrated BRS "
            "network. Glycaemic instability and HPA-axis dysregulation constrain "
            "mitochondrial recovery and bioenergetic reserve within BRS4 (Picard "
            "et al., 2014). Chronic stress activation propagates inflammatory "
            "signalling through BRS3 (Slavich & Irwin, 2014), while autonomic "
            "and cortisol dynamics directly modulate the conditions under which "
            "BRS1 sustains neurotransmitter regulation (Thayer et al., 2012). "
            "Gut–vagal signalling from BRS5 feeds back into autonomic and "
            "HPA-axis responsiveness, completing a peripheral–central loop "
            "through which dietary and lifestyle inputs influence whole-body "
            "load allocation.\n\n"
            "The framework makes BRS6→BRS1, BRS6→BRS3 and BRS6→BRS4 explicit "
            "as mutual constraints: when HPA-axis rhythm, glycaemic stability "
            "or autonomic recovery become chronically dysregulated, cumulative "
            "load progressively constrains performance across connected systems "
            "(McEwen, 2006). Landmark architectural evidence supports this "
            "organisation: Thayer et al. (2012) link heart-rate variability to "
            "prefrontal regulatory capacity; Picard et al. (2014) formalise "
            "mitochondrial allostatic load as the energetic cost of sustained "
            "stress-axis activation. BRS6 is therefore the principal gateway "
            "through which whole-body demand is translated into constraints on "
            "every connected regulatory system."
        ),
        "role_within_adaptive_network": (
            "BRS6 coordinates the resource-allocation logic through which "
            "whole-body energy, inflammatory burden and recovery capacity are "
            "proportionately distributed across the integrated network — the "
            "architectural tier most directly aligned with allostatic theory. "
            "When HPA-axis rhythm, glycaemic stability or autonomic recovery "
            "become chronically dysregulated, cumulative load progressively "
            "constrains performance across connected systems (McEwen, 2006). "
            "Meal timing, glycaemic balance and stress-recovery practices "
            "therefore become nutritionally meaningful because they influence "
            "the systemic logic through which adaptive regulatory capacities "
            "are mobilised and restored — not because they correct isolated "
            "hormone values, but because they support the resource-allocation "
            "architecture on which every other Biological Regulatory System "
            "depends."
        ),
        "references": {
            "adhd_landmark": [
                "Chang, J. P.-C., et al. (2021). Cortisol and inflammatory biomarker levels in youths with ADHD. Translational Psychiatry, 11, 388. https://doi.org/10.1038/s41398-021-01388-3",
                "Chang, J. P.-C., et al. (2020). Cortisol, inflammatory biomarkers and neurotrophins in children and adolescents with ADHD in Taiwan. Brain, Behavior, and Immunity, 88, 105–113. https://doi.org/10.1016/j.bbi.2020.05.017",
                "Lane, S. J., Reynolds, S., & Thacker, L. (2010). Sensory over-responsivity and ADHD differentiated using cortisol and autonomic markers. American Journal of Occupational Therapy, 64(3), 393–403. https://doi.org/10.5014/ajot.2010.09031",
                "Zametkin, A. J., et al. (1990). Cerebral glucose metabolism in adults with hyperactivity of childhood onset. New England Journal of Medicine, 323(20), 1361–1366. https://doi.org/10.1056/NEJM199011153232001",
                "Di Girolamo, G., et al. (2022). Prevalence of metabolic syndrome and insulin resistance in adult ADHD outpatients. Frontiers in Psychiatry, 13, 1078932. https://doi.org/10.3389/fpsyt.2022.1078932",
                "Marcelli, V., et al. (2025). Bridging ADHD and metabolic disorders: Shared mechanisms and clinical implications. Brain Sciences, 15(1), 42. https://doi.org/10.3390/brainsci15010042",
            ],
            "systems_integration": [
                "Thayer, J. F., Ahs, F., Fredrikson, M., Sollers, J. J., & Wager, T. D. (2012). A meta-analysis of heart rate variability and neuroimaging studies: Implications for heart rate variability as a marker of stress and health. Neuroscience & Biobehavioral Reviews, 36(2), 747–756. https://doi.org/10.1016/j.neubiorev.2011.11.009",
                "Picard, M., Juster, R.-P., & McEwen, B. S. (2014). Mitochondrial allostatic load puts the 'gluc' back in glucocorticoids. Nature Reviews Endocrinology, 10(5), 303–310. https://doi.org/10.1038/nrendo.2014.22",
                "Slavich, G. M., & Irwin, M. R. (2014). From stress to inflammation and major depressive disorder: A social signal transduction theory of depression. Psychological Bulletin, 140(3), 774–815. https://doi.org/10.1037/a0035302",
            ],
            "allostasis_and_resilience": [
                "McEwen, B. S. (1998). Protective and damaging effects of stress mediators. New England Journal of Medicine, 338(3), 171–179. https://doi.org/10.1056/NEJM199801083380307",
                "McEwen, B. S. (2006). Protective and damaging effects of stress mediators: Central role of the brain. Dialogues in Clinical Neuroscience, 8(4), 367–381. https://doi.org/10.31887/DCNS.2006.8.4/bmcewen",
                "Sterling, P., & Eyer, J. (1988). Allostasis: A new paradigm to explain arousal pathology. In S. Fisher & J. Reason (Eds.), Handbook of Life Stress, Cognition and Health. Wiley.",
            ],
        },
    },
]


def add_body_paragraphs(doc, text, heading_level=None, heading_size=14):
    if heading_level is not None:
        heading = doc.add_heading(text, level=heading_level)
        heading.runs[0].font.size = Pt(heading_size)
        return

    for para_text in text.split("\n\n"):
        p = doc.add_paragraph(para_text)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_numbered_references(doc, references):
    for ref in references:
        p = doc.add_paragraph(ref, style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def count_words(*texts):
    total = 0
    for text in texts:
        total += len(text.split())
    return total


def add_section(doc, section):
    add_body_paragraphs(doc, section["title"], heading_level=1, heading_size=14)

    add_body_paragraphs(
        doc, "Biological Regulatory Identity", heading_level=2, heading_size=13
    )
    add_body_paragraphs(doc, section["biological_regulatory_identity"])

    add_body_paragraphs(
        doc, "Landmark ADHD Evidence", heading_level=2, heading_size=13
    )
    add_body_paragraphs(doc, section["landmark_adhd_evidence"])

    add_body_paragraphs(doc, "Systems Integration", heading_level=2, heading_size=13)
    add_body_paragraphs(doc, section["systems_integration"])

    add_body_paragraphs(
        doc, "Role within the Adaptive Network", heading_level=2, heading_size=13
    )
    add_body_paragraphs(doc, section["role_within_adaptive_network"])

    doc.add_heading("References", level=2)

    refs = section["references"]
    add_body_paragraphs(doc, "ADHD landmark evidence", heading_level=3, heading_size=12)
    add_numbered_references(doc, refs["adhd_landmark"])

    add_body_paragraphs(doc, "Systems integration", heading_level=3, heading_size=12)
    add_numbered_references(doc, refs["systems_integration"])

    if refs["allostasis_and_resilience"]:
        add_body_paragraphs(
            doc, "Allostasis and resilience", heading_level=3, heading_size=12
        )
        add_numbered_references(doc, refs["allostasis_and_resilience"])

    return count_words(
        section["biological_regulatory_identity"],
        section["landmark_adhd_evidence"],
        section["systems_integration"],
        section["role_within_adaptive_network"],
    )


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.add_heading(
        "BRAIN Framework: Biological Regulatory System Summaries",
        level=0,
    )

    intro = doc.add_paragraph(DOCUMENT_INTRO)
    intro.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    intro.paragraph_format.space_after = Pt(12)
    for run in intro.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_page_break()
    add_body_paragraphs(
        doc,
        "Allostasis and the BRAIN Framework",
        heading_level=1,
        heading_size=14,
    )
    add_body_paragraphs(doc, ALLOSTASIS_FRAMEWORK_DISCUSSION)
    add_body_paragraphs(doc, "References", heading_level=2, heading_size=13)
    add_numbered_references(doc, ALLOSTASIS_FRAMEWORK_REFERENCES)
    allostasis_words = len(ALLOSTASIS_FRAMEWORK_DISCUSSION.split())
    print(f"Allostasis conceptual section word count: {allostasis_words}")

    counts = []
    for section in SECTIONS:
        doc.add_page_break()
        counts.append((section["title"].split("—")[0].strip(), add_section(doc, section)))

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)

    print(f"Saved: {OUTPUT}")
    for label, count in counts:
        print(f"{label} body word count: {count}")


if __name__ == "__main__":
    main()
