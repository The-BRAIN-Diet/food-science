#!/usr/bin/env python3
"""Generate BRS4 structure + ADHD studies manuscript summary as Word document."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

OUTPUT = (
    "/Users/paulhouston/Food Science/food-science/manuscript/"
    "BRS4-ADHD-structure-and-evidence-summary.docx"
)

TITLE = (
    "BRS4 — Mitochondrial Function & Bioenergetics: "
    "Framework Structure and ADHD-Relevant Evidence"
)

BODY = (
    "Within the BRAIN Framework, BRS4 (Mitochondrial Function & Bioenergetics) describes the "
    "integrated cellular capacity that sustains ATP production, fuel flexibility, mitochondrial "
    "protection, and longer-term mitochondrial quality control. Four Functional Mechanisms integrate "
    "nine Primary Mechanisms: BRS4(FM1) Cellular Bioenergetics (PM1–PM3: electron transport chain "
    "function, NAD⁺ metabolism, creatine–phosphocreatine buffering); BRS4(FM2) Mitochondrial "
    "Resilience & Redox Stability (PM4–PM5: ROS production and control, mitochondrial protection "
    "and redox integrity); BRS4(FM3) Substrate Utilisation & Flexibility (PM6–PM8: carnitine-mediated "
    "fat transport, ketone utilisation, metabolic fuel switching); and BRS4(FM4) Mitochondrial Capacity "
    "Expansion & Adaptation (PM9: mitochondrial biogenesis). Key Constraints include BRS4(KC1) "
    "Macronutrient Substrate Availability and BRS4(KC2) Mitochondrial Cofactor Sufficiency — shared "
    "pools constraining respiratory-chain function, substrate delivery, and redox buffering. BRS4 "
    "cross-links to BRS3 antioxidant and inflammatory context and to BRS5 gut-derived metabolite "
    "biology where butyrate and related SCFAs influence brain energy metabolism.\n\n"
    "In ADHD, mitochondrial biology is interpreted as a modifiable regulatory system rather than "
    "a single universal deficit. Direct ADHD evidence includes a controlled carnitine supplementation "
    "trial reporting behavioural and functional improvements in children with ADHD (van Oudheusden & "
    "Scholte, 2002); paediatric case–control work showing elevated glutathione and oxidative-stress "
    "markers with implications for mitochondrial lactate metabolism and ROS neutralisation (Verlaet "
    "et al., 2019); cybrid models from ADHD platelets demonstrating reduced cellular and mitochondrial "
    "respiration, lower complex V activity, membrane-potential loss, and elevated oxidative stress "
    "(Verma et al., 2016); and narrative reviews synthesising ADHD-relevant mitochondrial biomarker, "
    "oxidative-stress, and genetic-variation literature including mtDNA copy number and haplogroup "
    "associations (Öğütlü et al., 2022) and mitophagy-related bioenergetic regulation (Almutairi "
    "et al., 2024). Inferred mechanistic evidence extends from general mitochondrial and nutrition "
    "science: micronutrient cofactor sufficiency for respiratory-chain and brain-relevant pathways "
    "(Tardy et al., 2020); CoQ10 roles in electron transport and neuronal antioxidant protection "
    "(Crane, 2001); creatine-supported ATP recycling in neurons (Avgerinos et al., 2018); polyphenol-"
    "derived urolithin A and mitophagy pathways in non-ADHD intervention contexts (Singh et al., "
    "2022; Andreux et al., 2019; Hou et al., 2024); and butyrate-related enhancement of mitochondrial "
    "function under oxidative stress, interpreted through gut–brain metabolic contexts (Rose et al., "
    "2018). Human dietary intervention studies directly measuring mitochondrial functional mechanisms "
    "in ADHD remain limited; dedicated BRS4(SM-SNP) pages for mtDNA and polymorphism associations "
    "are pending framework expansion.\n\n"
    "Collectively, these findings do not imply that mitochondrial dysfunction is universal in ADHD, "
    "nor that individual biomarkers define the disorder. Instead, they support the BRAIN Framework "
    "proposition that ADHD-relevant bioenergetic phenotypes emerge from interacting constraints on "
    "electron transport, substrate delivery, redox stability, phosphocreatine buffering, and "
    "mitochondrial quality control — providing multiple diet-actionable entry points beyond isolated "
    "mitochondrial supplement models."
)

REFERENCES = [
    "Almutairi, M. M., Althekair, A., Almutairi, F., Alatabani, M., & Alsaikhan, A. (2024). Mitochondrial dysfunction and mitophagy in ADHD: Cellular and molecular mechanisms. Saudi Pharmaceutical Journal, 32(12). https://doi.org/10.1016/j.jsps.2024.102212",
    "Andreux, P. A., Blanco-Bose, W., Ryu, D., et al. (2019). The mitophagy activator urolithin A is safe and induces a molecular signature of improved mitochondrial and cellular health in humans. Nature Metabolism, 1(6). https://doi.org/10.1038/s42255-019-0073-4",
    "Avgerinos, K. I., Spyrou, N., Bougioukas, K. I., & Kapogiannis, D. (2018). Effects of creatine supplementation on cognitive function of healthy individuals: A systematic review of randomized controlled trials. Clinical Nutrition, 37(6), 2184–2190. https://doi.org/10.1016/j.clnu.2017.12.002",
    "Crane, F. L. (2001). Biochemical functions of coenzyme Q10. Journal of the American College of Nutrition, 20(6), 591–598. https://doi.org/10.1080/07315724.2001.10719063",
    "Hou, Y., Chu, X., Park, J.-H., et al. (2024). Urolithin A improves Alzheimer's disease cognition and restores mitophagy and lysosomal functions. Alzheimer's & Dementia, 20(6), 4212–4233. https://doi.org/10.1002/alz.13847",
    "Öğütlü, H., Kasak, M., & Tabur, S. T. (2022). Mitochondrial dysfunction in attention deficit hyperactivity disorder. The Eurasian Journal of Medicine, 54(Suppl 1), S187–S195. https://doi.org/10.5152/eurasianjmed.2022.22187",
    "Rose, S., Bennuri, S. C., Davis, J. E., et al. (2018). Butyrate enhances mitochondrial function during oxidative stress in cell lines from boys with autism. Translational Psychiatry, 8(1). https://doi.org/10.1038/s41398-017-0089-z",
    "Singh, A., D'Amico, D., Andreux, P. A., et al. (2022). Direct supplementation with urolithin A overcomes limitations of dietary exposure and gut microbiome variability in healthy adults. European Journal of Clinical Nutrition, 76(2), 297–308. https://doi.org/10.1038/s41430-021-00950-1",
    "Tardy, A.-L., Pouteau, E., Marquez, D., Yilmaz, C., & Scholey, A. (2020). Vitamins and minerals for energy, fatigue and cognition: A narrative review of the biochemical and clinical evidence. Nutrients, 12(1), 228. https://doi.org/10.3390/nu12010228",
    "van Oudheusden, L. J., & Scholte, H. R. (2002). Efficacy of carnitine in the treatment of children with attention-deficit hyperactivity disorder. Prostaglandins, Leukotrienes and Essential Fatty Acids, 67(1), 33–38. https://doi.org/10.1054/plef.2002.0378",
    "Verlaet, A. A. J., Breynaert, A., Ceulemans, B., et al. (2019). Oxidative stress and immune aberrancies in attention-deficit/hyperactivity disorder (ADHD): A case–control comparison. European Child & Adolescent Psychiatry, 28(5), 719–729. https://doi.org/10.1007/s00787-018-1239-4",
    "Verma, P., Singh, A., Nthenge-Ngumbau, D. N., et al. (2016). Attention deficit-hyperactivity disorder suffers from mitochondrial dysfunction. BBA Clinical, 6, 153–158. https://doi.org/10.1016/j.bbacli.2016.10.003",
]


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading(TITLE, level=1)
    title.runs[0].font.size = Pt(14)

    for para_text in BODY.split("\n\n"):
        p = doc.add_paragraph(para_text)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    doc.add_heading("References", level=2)
    for ref in REFERENCES:
        p = doc.add_paragraph(ref, style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    import os

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)

    word_count = len(BODY.split())
    print(f"Saved: {OUTPUT}")
    print(f"Body word count: {word_count}")


if __name__ == "__main__":
    main()
